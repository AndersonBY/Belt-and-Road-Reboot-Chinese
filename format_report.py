# -*- coding: utf-8 -*-
# @Author: Bi Ying
# @Date:   2023-11-10 12:51:56
# @Last Modified by:   Bi Ying
# @Last Modified time: 2023-11-11 18:22:24
import re
import os
import shutil
import base64
from io import BytesIO
from pathlib import Path

import mammoth
from PIL import Image
from docx import Document
from docx.table import Table


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def is_partial_end(text: str):
    text = text.strip()
    if len(text) == 0:
        return True
    return text[-1] not in [".", "!", "?", "\n", "”", '"']


def is_partial_start(text: str):
    text = text.strip()
    if len(text) == 0:
        return True
    return text[0].islower() or text[0] in ["(", "$"]


def table_to_markdown(table):
    markdown_str = ""

    for cell in table.rows[0].cells:
        markdown_str += f"| {cell.text} "
    markdown_str += "|\n"

    markdown_str += "|---" * len(table.columns) + "|\n"

    for row in table.rows[1:]:
        for cell in row.cells:
            markdown_str += f"| {cell.text} "
        markdown_str += "|\n"

    return markdown_str


def replace_base64_image_with_local_image(
    chapter_index: int, markdown_text: str, image_folder: Path, output_folder: Path
):
    # 匹配base64图片的正则表达式
    pattern = r"!\[\]\((data:image\/(\w+);base64,([a-zA-Z0-9+/=]+))\)"

    # 对每一个匹配项进行处理
    for image_index, match in enumerate(re.findall(pattern, markdown_text)):
        data_uri, image_format, base64_data = match

        # 解码base64图片
        image_data = base64.b64decode(base64_data)
        image = Image.open(BytesIO(image_data))

        # 生成UUID文件名并保存图片
        filename = f"{chapter_index}-{image_index}.{image_format}"
        filepath = image_folder / filename
        image.save(filepath)

        # 替换Markdown中的图片链接
        markdown_text = markdown_text.replace(data_uri, str(os.path.relpath(filepath, output_folder)))

    return markdown_text


def format_report(chapter_index: int, file: str | Path, output_folder: Path, image_folder: Path):
    document = Document(file)
    reference_document: Document = Document()

    annotation_texts = []
    annotation_paragraphs = []
    normal_paragraphs = []
    reference_paragraphs = []

    for paragraph in document.iter_inner_content():
        if isinstance(paragraph, Table):
            # TODO: 表格插入到文档
            # table_md_text = table_to_markdown(paragraph)
            continue

        if paragraph.text.strip().lower() == "references":
            reference_paragraphs.append(paragraph)
            continue
        if len(reference_paragraphs) > 0:
            reference_paragraphs.append(paragraph)
            continue
        if len(paragraph.runs) == 0:
            continue
        if len(paragraph.runs) > 0 and paragraph.runs[0].font.superscript:
            annotation_paragraphs.append(paragraph)
            annotation_texts.append(paragraph.text)
        elif paragraph.style.name == "Body Text" and len(paragraph.text) > 0:
            normal_paragraphs.append(paragraph)
        elif len(paragraph.runs) > 0 and paragraph.runs[0].font.italic:
            normal_paragraphs.append(paragraph)
        elif paragraph.style.name == "Normal" and not paragraph.text.startswith("Notes:"):
            # 对于不同转换软件生成的文档，这里的字体大小可能不同
            if paragraph.runs[0].font.size == 114300:
                annotation_paragraphs.append(paragraph)
                annotation_texts.append(paragraph.text)
            else:
                normal_paragraphs.append(paragraph)
        else:
            normal_paragraphs.append(paragraph)

    for paragraph in annotation_paragraphs:
        delete_paragraph(paragraph)

    for paragraph in reference_paragraphs:
        reference_document.add_paragraph(paragraph.text)
        inserted_p = reference_document._body._body._insert_p(paragraph._p)
        if paragraph._p.get_or_add_pPr().numPr:
            inserted_p.style = "ListNumber"
        delete_paragraph(paragraph)

    # 遍历所有normal_paragraphs
    # 找出上一段未结束接续的，需要和上一段合并并删掉这一段
    partial_paragraphs = []
    for index, paragraph in enumerate(normal_paragraphs):
        for run in paragraph.runs:
            if run.font.superscript:
                run.text = f"[{run.text}]"
        if index == 0:
            continue
        if len(paragraph.text) == 0:
            continue

        # 如果该段开头是小写，或者上一段结尾不是标点符号加换行，说明是上一段未结束接续的
        if is_partial_start(paragraph.text) and is_partial_end(normal_paragraphs[index - 1].text):
            # 上一段可能也是partial，所以要判断一下，找到真正开始的paragraph
            paragraph_start_index = index - 1
            while (
                paragraph_start_index >= 0
                and is_partial_start(normal_paragraphs[paragraph_start_index].text)
                and is_partial_end(normal_paragraphs[paragraph_start_index - 1].text)
            ):
                paragraph_start_index -= 1

            if (
                len(normal_paragraphs[paragraph_start_index].text) == 0
                or normal_paragraphs[paragraph_start_index].text[-1] == "\n"
            ):
                normal_paragraphs[paragraph_start_index].text = normal_paragraphs[paragraph_start_index].text[:-1]
            normal_paragraphs[paragraph_start_index].text += f" {paragraph.text}"
            partial_paragraphs.append(paragraph)

    for paragraph in partial_paragraphs:
        delete_paragraph(paragraph)

    document.add_heading("# Annotation", 0)
    annotation_texts = filter(lambda x: len(x) > 0, annotation_texts)
    document.add_paragraph("\n\n".join(annotation_texts))

    document.save(output_folder / f"[Formatted]{file.stem}.docx")
    reference_document.save(output_folder / f"[References]{file.stem}.docx")

    with open(output_folder / f"[Formatted]{file.stem}.docx", "rb") as docx_file:
        docx_data = mammoth.convert_to_markdown(docx_file)
        markdown_text = replace_base64_image_with_local_image(
            chapter_index,
            docx_data.value,
            image_folder,
            output_folder,
        )
        markdown_text = f"> 格式化整理及翻译：Maker毕\n\n{markdown_text}"

    with open(output_folder / f"{file.stem}.md", "w", encoding="utf8") as text_file:
        text_file.write(markdown_text)


image_folder = Path("images")
image_folder.mkdir(exist_ok=True)

input_folder = Path("input")
output_folder = Path("formatted")
for i in range(1, 5):
    file = input_folder / Path(f"Belt_and_Road_Reboot_Chapter_{i}.docx")
    format_report(i, file, output_folder, image_folder)

proofread_folder = Path("proofread")
proofread_folder.mkdir(exist_ok=True)
Proofread_formatted_folder = proofread_folder / "formatted"
Proofread_formatted_folder.mkdir(exist_ok=True)
for file in output_folder.glob("*.md"):
    if Proofread_formatted_folder.joinpath(file.name).exists():
        print(f"File {file.name} already exists in {Proofread_formatted_folder}")
    else:
        shutil.copy2(file, Proofread_formatted_folder)
